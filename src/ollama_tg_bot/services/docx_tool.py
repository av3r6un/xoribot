from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DOCX_BLOCK_RE = re.compile(r'```(?:xoridocx|docx-json)\s*(.*?)```', re.S | re.I)
DOCX_BLOCK_START_RE = re.compile(r'```(?:xoridocx|docx-json)\b', re.I)
DOCX_BLOCK_PARTIAL_START_RE = re.compile(r'```(?:x|xo|xor|xori|xorid|xorido|xoridoc|xoridocx|d|do|doc|docx|docx-|docx-j|docx-js|docx-jso|docx-json)?$', re.I)

DOCX_TOOL_PROMPT = '''
Создай Word/.docx документ.
Не отвечай обычным текстом. Верни только один fenced-блок ```xoridocx
с JSON-разметкой документа. Разметку документа создаёшь ты сразу,
бот только запечатывает её в .docx.

Формат:
```xoridocx
{
  "filename": "document.docx",
  "title": "Название документа",
  "properties": {"author": "XoriBot", "subject": "Тема"},
  "blocks": [
    {"type": "heading", "level": 1, "text": "Заголовок"},
    {"type": "paragraph", "runs": [
      {"text": "Обычный текст "},
      {"text": "жирный", "bold": true},
      {"text": " и курсив", "italic": true}
    ]},
    {"type": "list", "ordered": false, "items": ["Пункт 1", "Пункт 2"]},
    {"type": "table", "headers": ["Колонка 1", "Колонка 2"], "rows": [["A", "B"]]},
    {"type": "page_break"}
  ]
}
```

Поддерживаемые block.type: heading, paragraph, list, table, page_break.
Для paragraph можно использовать text или runs. Для runs доступны text, bold, italic, underline.
Для heading level должен быть 1-4. Filename должен заканчиваться на .docx.
'''.strip()

DOCX_DELEGATE_PROMPT = '''
Пользователь попросил создать Word/.docx документ.
Не отвечай обычным текстом. Верни только строку @docx и один fenced-блок ```xoridocx
с готовой JSON-разметкой документа. Разметку документа создаёшь ты сразу.
Не добавляй пояснения до или после блока.

Формат:
@docx
```xoridocx
{
  "filename": "document.docx",
  "title": "Название документа",
  "properties": {"author": "XoriBot", "subject": "Тема"},
  "blocks": [
    {"type": "heading", "level": 1, "text": "Заголовок"},
    {"type": "paragraph", "runs": [
      {"text": "Обычный текст "},
      {"text": "жирный", "bold": true},
      {"text": " и курсив", "italic": true}
    ]},
    {"type": "list", "ordered": false, "items": ["Пункт 1", "Пункт 2"]},
    {"type": "table", "headers": ["Колонка 1", "Колонка 2"], "rows": [["A", "B"]]}
  ]
}
```
'''.strip()


class DocxToolError(Exception):
  user_message = 'Не удалось создать .docx. Проверь разметку документа.'

  def __init__(self, message: str | None = None, user_message: str | None = None) -> None:
    super().__init__(message or self.user_message)
    if user_message: self.user_message = user_message


def extract_docx_specs(text: str) -> tuple[str, list[dict]]:
  specs: list[dict] = []

  def replace(match: re.Match) -> str:
    raw = match.group(1).strip()
    if raw:
      specs.append(_load_spec(raw))
    return ''

  visible_text = _strip_docx_tag(DOCX_BLOCK_RE.sub(replace, text)).strip()
  return visible_text, specs


def visible_docx_text(text: str) -> str:
  visible = DOCX_BLOCK_RE.sub('', text)
  start = DOCX_BLOCK_START_RE.search(visible)
  if start: visible = visible[:start.start()]
  partial_start = DOCX_BLOCK_PARTIAL_START_RE.search(visible)
  if partial_start: visible = visible[:partial_start.start()]
  return _strip_docx_tag(visible).strip()


def build_docx(spec: dict, destination: Path) -> Path:
  filename = _safe_filename(spec.get('filename') or spec.get('title') or 'document.docx')
  path = destination / filename

  document = Document()
  _apply_properties(document, spec)

  title = _clean_text(spec.get('title'))
  if title:
    document.add_heading(title, level=0)

  blocks = spec.get('blocks') or spec.get('sections') or []
  if not isinstance(blocks, list):
    raise DocxToolError('blocks must be a list')

  for block in blocks:
    if not isinstance(block, dict): continue
    _add_block(document, block)

  if not blocks and not title:
    raise DocxToolError('empty document', user_message='Модель вернула пустой документ.')

  document.save(path)
  return path


def _load_spec(raw: str) -> dict:
  try:
    data = json.loads(raw)
  except json.JSONDecodeError as exc:
    raise DocxToolError('invalid docx json') from exc
  if not isinstance(data, dict): raise DocxToolError('docx json must be an object')
  return data


def _apply_properties(document: Document, spec: dict) -> None:
  properties = spec.get('properties') or {}
  if not isinstance(properties, dict): properties = {}

  title = _clean_text(spec.get('title'))
  if title: document.core_properties.title = title

  author = _clean_text(properties.get('author'))
  subject = _clean_text(properties.get('subject'))
  keywords = _clean_text(properties.get('keywords'))
  if author: document.core_properties.author = author
  if subject: document.core_properties.subject = subject
  if keywords: document.core_properties.keywords = keywords


def _add_block(document: Document, block: dict) -> None:
  block_type = str(block.get('type') or 'paragraph').strip().lower()
  if block_type == 'heading':
    level = _int_value(block.get('level'), 1)
    level = min(max(level, 1), 4)
    document.add_heading(_clean_text(block.get('text')), level=level)
    return

  if block_type == 'paragraph':
    paragraph = document.add_paragraph()
    _apply_alignment(paragraph, block.get('alignment'))
    _add_runs(paragraph, block)
    return

  if block_type == 'list':
    _add_list(document, block)
    return

  if block_type == 'table':
    _add_table(document, block)
    return

  if block_type == 'page_break':
    document.add_page_break()


def _add_list(document: Document, block: dict) -> None:
  items = block.get('items') or []
  if not isinstance(items, list): return
  style = 'List Number' if block.get('ordered') else 'List Bullet'
  for item in items:
    paragraph = document.add_paragraph(style=style)
    if isinstance(item, dict):
      _add_runs(paragraph, item)
    else:
      paragraph.add_run(_clean_text(item))


def _add_table(document: Document, block: dict) -> None:
  headers = block.get('headers') or []
  rows = block.get('rows') or []
  if not isinstance(headers, list): headers = []
  if not isinstance(rows, list): rows = []
  width = max(len(headers), _max_row_width(rows))
  if width <= 0: return

  table_rows = len(rows) + (1 if headers else 0)
  table = document.add_table(rows=table_rows, cols=width)
  table.style = 'Table Grid'

  row_offset = 0
  if headers:
    for index, value in enumerate(headers[:width]):
      cell = table.cell(0, index)
      cell.text = _clean_text(value)
      for paragraph in cell.paragraphs:
        for run in paragraph.runs:
          run.bold = True
    row_offset = 1

  for row_index, row in enumerate(rows):
    if not isinstance(row, list): row = [row]
    for col_index, value in enumerate(row[:width]):
      table.cell(row_index + row_offset, col_index).text = _clean_text(value)


def _add_runs(paragraph: Any, payload: dict) -> None:
  runs = payload.get('runs')
  if not isinstance(runs, list):
    text = _clean_text(payload.get('text'))
    if text: paragraph.add_run(text)
    return

  for item in runs:
    if isinstance(item, dict):
      run = paragraph.add_run(_clean_text(item.get('text')))
      run.bold = bool(item.get('bold'))
      run.italic = bool(item.get('italic'))
      run.underline = bool(item.get('underline'))
      size = _int_value(item.get('size'), 0)
      if size > 0: run.font.size = Pt(size)
    else:
      paragraph.add_run(_clean_text(item))


def _apply_alignment(paragraph: Any, value: Any) -> None:
  alignment = str(value or '').strip().lower()
  if alignment == 'center':
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  elif alignment == 'right':
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  elif alignment == 'justify':
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _safe_filename(value: Any) -> str:
  filename = _clean_text(value).strip() or 'document.docx'
  filename = re.sub(r'[\\/:\*\?"<>\|]+', '-', filename)
  filename = filename[:100].strip('. ') or 'document.docx'
  if not filename.lower().endswith('.docx'): filename = f'{filename}.docx'
  return filename


def _max_row_width(rows: list) -> int:
  width = 0
  for row in rows:
    if isinstance(row, list):
      width = max(width, len(row))
    else:
      width = max(width, 1)
  return width


def _int_value(value: Any, default: int) -> int:
  try:
    return int(value)
  except (TypeError, ValueError):
    return default


def _clean_text(value: Any) -> str:
  if value is None: return ''
  return str(value).replace('\x00', '').strip()


def _strip_docx_tag(text: str) -> str:
  return re.sub(r'^\s*@docx\b[:\s-]*', '', text, flags=re.I)
